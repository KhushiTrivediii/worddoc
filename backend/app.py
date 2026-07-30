import os
import json
import tempfile
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

app = Flask(__name__, static_folder='../frontend', static_url_path='')
CORS(app)  # Enable CORS for frontend communication

@app.route('/')
def index():
    return send_file(os.path.join(app.static_folder, 'index.html'))

def convert_doc_to_docx(doc_path):
    """Converts a .doc file to .docx using Microsoft Word COM automation."""
    try:
        import win32com.client as win32
        word = win32.Dispatch('Word.Application')
        word.Visible = False
        
        abs_doc_path = os.path.abspath(doc_path)
        doc = word.Documents.Open(abs_doc_path)
        abs_docx_path = os.path.splitext(abs_doc_path)[0] + '.docx'
        
        # SaveAs2 with FileFormat=12 (docx format)
        doc.SaveAs2(abs_docx_path, FileFormat=12)
        doc.Close(False)
        word.Quit()
        return abs_docx_path
    except Exception as e:
        print(f"Word COM conversion failed: {e}")
        raise e

# Configuration
UPLOAD_FOLDER = tempfile.gettempdir()
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def set_cell_background(cell, fill_hex):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set cell padding (in dxas: 20 dxa = 1 pt)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(margin)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_code_block(doc, code_text):
    """Adds a code block using a single-cell table with light-gray shading."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell = table.cell(0, 0)
    set_cell_background(cell, "F2F2F2")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    # Remove borders
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="single" w:sz="12" w:space="0" w:color="CCCCCC"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def apply_text_formatting(run, font_name='Arial', size_pt=11, bold=False, italic=False, color_rgb=None):
    """Applies standard font and color properties to a text run."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

@app.route('/api/health', methods=['GET'])
def health():
    return jsonify({"status": "healthy", "message": "Word Generator Backend is running!"})

@app.route('/api/analyze-template', methods=['POST'])
def analyze_template():
    if 'template' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['template']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    if not (file.filename.endswith('.docx') or file.filename.endswith('.doc')):
        return jsonify({"error": "File must be a .doc or .docx document"}), 400
    
    # Save template to a temporary location
    is_doc = file.filename.endswith('.doc')
    temp_suffix = '.doc' if is_doc else '.docx'
    temp_path = os.path.join(app.config['UPLOAD_FOLDER'], tempfile.mktemp(suffix=temp_suffix))
    file.save(temp_path)
    
    # If legacy .doc file, convert to .docx first
    if is_doc:
        try:
            temp_docx_path = convert_doc_to_docx(temp_path)
            os.remove(temp_path)
            temp_path = temp_docx_path
        except Exception as e:
            if os.path.exists(temp_path):
                os.remove(temp_path)
            return jsonify({"error": f"Failed to convert legacy .doc file to .docx: {str(e)}"}), 500

    
    try:
        doc = DocxTemplate(temp_path)
        variables = doc.get_undeclared_template_variables()
        
        # Categorize variables for rich UI experience
        categorized_variables = []
        for var in sorted(list(variables)):
            lowered = var.lower()
            # If the variable name hints at being an image/screenshot
            is_img = any(x in lowered for x in ['image', 'img', 'pic', 'photo', 'screenshot', 'paste'])
            categorized_variables.append({
                "name": var,
                "type": "image" if is_img else "text",
                "label": var.replace('_', ' ').replace('-', ' ').title()
            })
            
        return jsonify({
            "filename": file.filename,
            "variables": categorized_variables
        })
    except Exception as e:
        return jsonify({"error": f"Failed to parse template: {str(e)}"}), 500
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)

@app.route('/api/generate-from-template', methods=['POST'])
def generate_from_template():
    if 'template' not in request.files:
        return jsonify({"error": "No template file provided"}), 400
        
    template_file = request.files['template']
    
    # Save the template to a temporary location
    is_doc = template_file.filename.endswith('.doc')
    temp_suffix = '.doc' if is_doc else '.docx'
    temp_template_path = os.path.join(app.config['UPLOAD_FOLDER'], tempfile.mktemp(suffix=temp_suffix))
    template_file.save(temp_template_path)
    
    # If legacy .doc file, convert to .docx first
    if is_doc:
        try:
            temp_docx_path = convert_doc_to_docx(temp_template_path)
            os.remove(temp_template_path)
            temp_template_path = temp_docx_path
        except Exception as e:
            if os.path.exists(temp_template_path):
                os.remove(temp_template_path)
            return jsonify({"error": f"Failed to convert legacy .doc file to .docx: {str(e)}"}), 500

    
    # Parse text context variables
    context_data = request.form.get('context', '{}')
    try:
        context = json.loads(context_data)
    except Exception:
        return jsonify({"error": "Invalid json context data"}), 400
    
    # Keep track of temporary image files to clean them up later
    temp_image_files = []
    
    try:
        doc = DocxTemplate(temp_template_path)
        
        # Process files (screenshots/images)
        for key in request.files:
            if key == 'template':
                continue
                
            img_file = request.files[key]
            if img_file and img_file.filename != '':
                # Save to a temporary file
                temp_img_path = os.path.join(app.config['UPLOAD_FOLDER'], tempfile.mktemp(suffix='.png'))
                img_file.save(temp_img_path)
                temp_image_files.append(temp_img_path)
                
                # Create InlineImage or Subdoc and bind it to the context
                # Default width is 6.0 inches (fits standard letter page margins nicely)
                # If custom width is passed in the context metadata, use that
                custom_width = context.get(f"{key}_width", 6.0)
                try:
                    custom_width = float(custom_width)
                except ValueError:
                    custom_width = 6.0
                    
                # If instruction text is provided for this screenshot, compile as subdocument
                instruction_text = context.get(f"{key}_instruction", "").strip()
                if instruction_text:
                    subdoc = doc.new_subdoc()
                    
                    # Image paragraph (centered)
                    p_img = subdoc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(8)
                    p_img.paragraph_format.space_after = Pt(4)
                    run_img = p_img.add_run()
                    run_img.add_picture(temp_img_path, width=Inches(custom_width))
                    
                    # Caption/Instruction text paragraph (centered, gray, italic)
                    p_text = subdoc.add_paragraph()
                    p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_text.paragraph_format.space_before = Pt(2)
                    p_text.paragraph_format.space_after = Pt(12)
                    run_text = p_text.add_run(instruction_text)
                    apply_text_formatting(run_text, size_pt=9.5, italic=True, color_rgb=RGBColor(0x47, 0x55, 0x69))
                    
                    context[key] = subdoc
                else:
                    context[key] = InlineImage(doc, temp_img_path, width=Inches(custom_width))
                
        # Render the template
        doc.render(context)
        
        # Save output document
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], tempfile.mktemp(suffix='.docx'))
        doc.save(output_path)
        
        return send_file(
            output_path,
            as_attachment=True,
            download_name=f"Generated_{template_file.filename}",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return jsonify({"error": f"Failed to generate document: {str(e)}"}), 500
    finally:
        # Clean up files
        if os.path.exists(temp_template_path):
            os.remove(temp_template_path)
        for f in temp_image_files:
            if os.path.exists(f):
                os.remove(f)

@app.route('/api/generate-from-scratch', methods=['POST'])
def generate_from_scratch():
    document_data_str = request.form.get('document_data', '{}')
    try:
        doc_data = json.loads(document_data_str)
    except Exception:
        return jsonify({"error": "Invalid document data json"}), 400
        
    style = doc_data.get('style', 'standard')
    title_text = doc_data.get('title', 'Untitled Document')
    blocks = doc_data.get('blocks', [])
    
    # Check if a base template doc/docx is uploaded
    base_template_path = None
    if 'base_template' in request.files:
        base_template_file = request.files['base_template']
        if base_template_file and (base_template_file.filename.endswith('.docx') or base_template_file.filename.endswith('.doc')):
            is_doc_base = base_template_file.filename.endswith('.doc')
            temp_suffix = '.doc' if is_doc_base else '.docx'
            base_template_path = os.path.join(app.config['UPLOAD_FOLDER'], tempfile.mktemp(suffix=temp_suffix))
            base_template_file.save(base_template_path)
            
            # Convert legacy doc base template to docx
            if is_doc_base:
                try:
                    temp_docx_path = convert_doc_to_docx(base_template_path)
                    os.remove(base_template_path)
                    base_template_path = temp_docx_path
                except Exception as e:
                    if os.path.exists(base_template_path):
                        os.remove(base_template_path)
                    return jsonify({"error": f"Failed to convert base .doc template to .docx: {str(e)}"}), 500


    # Check if a title logo/banner image is uploaded
    title_logo_path = None
    if 'title_logo' in request.files:
        title_logo_file = request.files['title_logo']
        if title_logo_file and title_logo_file.filename != '':
            title_logo_path = os.path.join(app.config['UPLOAD_FOLDER'], tempfile.mktemp(suffix='.png'))
            title_logo_file.save(title_logo_path)

    # Save uploaded images to temp files and map them to their image_id
    image_map = {}
    temp_image_files = []
    
    for key in request.files:
        if key in ['base_template', 'title_logo']:
            continue
        img_file = request.files[key]
        if img_file and img_file.filename != '':
            temp_img_path = os.path.join(app.config['UPLOAD_FOLDER'], tempfile.mktemp(suffix='.png'))
            img_file.save(temp_img_path)
            temp_image_files.append(temp_img_path)
            image_map[key] = temp_img_path

            
    try:
        # Create a new Word document or open from base template
        if base_template_path:
            doc = Document(base_template_path)
        else:
            doc = Document()
            # Page Margins (1 inch all sides) only applied for blank documents,
            # since a template already has its own margins defined!
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)

            
        # Global Color Palettes & Fonts based on style
        font_name = 'Segoe UI' if style in ['screenshot-log', 'code-doc'] else 'Calibri'
        primary_color = RGBColor(0x0f, 0x17, 0x2a)  # Slate 900
        accent_color = RGBColor(0x4f, 0x46, 0xe5)   # Indigo 600
        secondary_color = RGBColor(0x47, 0x55, 0x69) # Slate 600
        
        if style == 'meeting-notes':
            font_name = 'Arial'
            accent_color = RGBColor(0x0f, 0x76, 0x6e)  # Teal 700
        elif style == 'code-doc':
            accent_color = RGBColor(0x02, 0x84, 0xc7)  # Sky 600
            
        # Add Title Logo/Banner if provided
        if title_logo_path:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER if style != 'meeting-notes' else WD_ALIGN_PARAGRAPH.LEFT
            p_logo.paragraph_format.space_before = Pt(12)
            p_logo.paragraph_format.space_after = Pt(12)
            run_logo = p_logo.add_run()
            run_logo.add_picture(title_logo_path, width=Inches(3.5))

        # Add Document Title
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER if style != 'meeting-notes' else WD_ALIGN_PARAGRAPH.LEFT
        p_title.paragraph_format.space_before = Pt(12)
        p_title.paragraph_format.space_after = Pt(24)
        
        run_title = p_title.add_run(title_text)
        apply_text_formatting(run_title, font_name=font_name, size_pt=24, bold=True, color_rgb=accent_color)
        
        # Style specific title structures
        if style == 'meeting-notes':
            # Create a nice metadata table for meetings
            table = doc.add_table(rows=2, cols=2)
            table.autofit = True
            
            # Label Cells
            cell_tl = table.cell(0, 0)
            cell_tr = table.cell(0, 1)
            cell_bl = table.cell(1, 0)
            cell_br = table.cell(1, 1)
            
            # Fetch variables from metadata if available
            meeting_date = doc_data.get('meeting_date', 'Date: Not specified')
            meeting_attendees = doc_data.get('meeting_attendees', 'Attendees: Not specified')
            meeting_facilitator = doc_data.get('meeting_facilitator', 'Facilitator: Not specified')
            meeting_location = doc_data.get('meeting_location', 'Location: Not specified')
            
            cell_tl.text = f"Date: {meeting_date}"
            cell_tr.text = f"Location: {meeting_location}"
            cell_bl.text = f"Attendees: {meeting_attendees}"
            cell_br.text = f"Facilitator: {meeting_facilitator}"
            
            for row in table.rows:
                for cell in row.cells:
                    set_cell_background(cell, "F8FAFC")
                    set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(2)
                        for run in paragraph.runs:
                            apply_text_formatting(run, font_name=font_name, size_pt=10, color_rgb=secondary_color)
            
            # Spacing after table
            doc.add_paragraph().paragraph_format.space_after = Pt(12)
            
        # Process Blocks
        step_number = 1
        for block in blocks:
            b_type = block.get('type')
            
            if b_type == 'heading':
                level = int(block.get('level', 1))
                p = doc.add_paragraph()
                
                # Indent heading numbering if it is a screenshot-log
                text = block.get('content', '')
                if style == 'screenshot-log' and level == 1:
                    text = f"Step {step_number}: {text}"
                    step_number += 1
                
                # Heading sizing and spacing
                if level == 1:
                    size = 18
                    color = accent_color
                    p.paragraph_format.space_before = Pt(16)
                    p.paragraph_format.space_after = Pt(6)
                elif level == 2:
                    size = 14
                    color = primary_color
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(4)
                else:
                    size = 12
                    color = secondary_color
                    p.paragraph_format.space_before = Pt(10)
                    p.paragraph_format.space_after = Pt(4)
                    
                run = p.add_run(text)
                apply_text_formatting(run, font_name=font_name, size_pt=size, bold=True, color_rgb=color)
                
            elif b_type == 'paragraph':
                text = block.get('content', '')
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                
                run = p.add_run(text)
                apply_text_formatting(run, font_name=font_name, size_pt=11, color_rgb=primary_color)
                
            elif b_type == 'code':
                text = block.get('content', '')
                add_code_block(doc, text)
                
            elif b_type == 'list':
                items = block.get('items', [])
                list_style = block.get('list_style', 'bullet')
                for item in items:
                    style_name = 'List Bullet' if list_style == 'bullet' else 'List Number'
                    p = doc.add_paragraph(style=style_name)
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.line_spacing = 1.15
                    
                    run = p.add_run(item)
                    apply_text_formatting(run, font_name=font_name, size_pt=11, color_rgb=primary_color)
                    
            elif b_type == 'image':
                img_id = block.get('image_id')
                caption = block.get('caption', '')
                width_in = float(block.get('width', 6.0))
                
                if img_id in image_map:
                    # Centered Image Paragraph
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(8)
                    p_img.paragraph_format.space_after = Pt(4)
                    
                    run_img = p_img.add_run()
                    run_img.add_picture(image_map[img_id], width=Inches(width_in))
                    
                    # Caption
                    if caption:
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_cap.paragraph_format.space_before = Pt(2)
                        p_cap.paragraph_format.space_after = Pt(12)
                        
                        run_cap = p_cap.add_run(caption)
                        apply_text_formatting(run_cap, font_name=font_name, size_pt=9.5, italic=True, color_rgb=secondary_color)
                        
        # Save output document
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], tempfile.mktemp(suffix='.docx'))
        doc.save(output_path)
        
        safe_title = "".join([c if c.isalnum() or c in [' ', '_', '-'] else '' for c in title_text])
        download_name = f"{safe_title.replace(' ', '_')}_{style}.docx"
        
        return send_file(
            output_path,
            as_attachment=True,
            download_name=download_name,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return jsonify({"error": f"Failed to generate scratch document: {str(e)}"}), 500
    finally:
        # Clean up images
        for f in temp_image_files:
            if os.path.exists(f):
                os.remove(f)
        # Clean up base template if it exists
        if base_template_path and os.path.exists(base_template_path):
            os.remove(base_template_path)
        # Clean up title logo if it exists
        if title_logo_path and os.path.exists(title_logo_path):
            os.remove(title_logo_path)



if __name__ == '__main__':
    # Start flask server on local port 5000
    app.run(port=5000, debug=True)
